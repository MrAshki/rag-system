"""Generate preprocessing test fixtures in several formats (TXT / DOCX / PDF).

Each fixture deliberately includes the failure cases we just fixed:
  - a junk first line "[]" and a "[Cover ...]" image-alt line (must be stripped),
  - ordinary paragraphs that START with "Part ..." / "Chapter 3 ..." (must stay body,
    not become section headings),
  - genuine short headings ("Chapter 1: ...", "فصل ۲: ...") that MUST be detected,
  - a bullet list and Persian + English text.

Run:  .\\venv\\Scripts\\python.exe tests\\preprocessing\\make_fixtures.py
"""
import os
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
FIXDIR = os.path.join(HERE, "fixtures")
os.makedirs(FIXDIR, exist_ok=True)


# ---------------------------------------------------------------- TXT
TXT = """[]

[Cover for The Sample Handbook, Author, Jane Doe]

The Sample Handbook

Chapter 1: Getting Started

This is the first real paragraph of content. It explains the basics in a couple of sentences so the chunker has a genuine block of body text to pack into a chunk.

Part of the difficulty in writing good tests is that ordinary paragraphs sometimes begin with words like Part or Chapter, and a naive heading detector will wrongly promote this entire long paragraph to a section heading even though it is plainly body text that runs on for a while.

Chapter 3 covered how the earlier material connects to this section, but this sentence is itself just a normal paragraph and must not be treated as a heading merely because it starts with the word Chapter.

فصل ۲: مبانی

این یک پاراگراف فارسی است که محتوای واقعی دارد و باید به‌عنوان متن عادی پردازش شود نه عنوان. هدف این است که مطمئن شویم متن فارسی به‌درستی و بدون آسیب حفظ می‌شود.

اهداف این بخش:
- مورد اول برای آزمایش فهرست
- مورد دوم برای آزمایش فهرست
- مورد سوم برای آزمایش فهرست
"""

with open(os.path.join(FIXDIR, "sample.txt"), "w", encoding="utf-8") as f:
    f.write(TXT)


# ---------------------------------------------------------------- DOCX
doc = Document()
doc.add_heading("The Report", level=1)
doc.add_paragraph("This is the introduction paragraph with real content explaining the purpose of the report in a sentence or two.")
doc.add_heading("Background", level=2)
doc.add_paragraph(
    "Part of the analysis showed that ordinary paragraphs which begin with the word Part "
    "should remain body text; this long sentence exists specifically to confirm the heading "
    "detector does not promote it to a heading."
)
doc.add_paragraph("این یک پاراگراف فارسی در یک فایل ورد است و باید سالم و به‌عنوان متن عادی پردازش شود.")
doc.add_paragraph("مورد اول", style="List Bullet")
doc.add_paragraph("مورد دوم", style="List Bullet")
doc.save(os.path.join(FIXDIR, "sample.docx"))


# ---------------------------------------------------------------- PDF (hand-built)
def _make_pdf(pages, path):
    """Minimal, dependency-free text PDF (Helvetica/ASCII). Each line is placed at an
    absolute position; blank entries add vertical space so pypdf sees paragraph breaks."""
    def esc(s):
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    n = len(pages)
    page_nums = [4 + i * 2 for i in range(n)]
    content_nums = [5 + i * 2 for i in range(n)]
    parts = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: ("<< /Type /Pages /Kids [" + " ".join(f"{p} 0 R" for p in page_nums)
            + f"] /Count {n} >>").encode(),
        3: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    for i, lines in enumerate(pages):
        ops = ["BT", "/F1 11 Tf"]
        y = 760
        for ln in lines:
            if ln == "":
                y -= 18
                continue
            ops.append(f"1 0 0 1 50 {y} Tm")
            ops.append(f"({esc(ln)}) Tj")
            y -= 16
        ops.append("ET")
        stream = "\n".join(ops).encode("latin-1")
        parts[content_nums[i]] = (b"<< /Length " + str(len(stream)).encode()
                                  + b" >>\nstream\n" + stream + b"\nendstream")
        parts[page_nums[i]] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_nums[i]} 0 R >>"
        ).encode()

    out = b"%PDF-1.4\n"
    off = {}
    mx = max(parts)
    for num in range(1, mx + 1):
        off[num] = len(out)
        out += f"{num} 0 obj\n".encode() + parts[num] + b"\nendobj\n"
    xref = len(out)
    out += f"xref\n0 {mx + 1}\n".encode() + b"0000000000 65535 f \n"
    for num in range(1, mx + 1):
        out += f"{off[num]:010d} 00000 n \n".encode()
    out += (b"trailer\n" + f"<< /Size {mx + 1} /Root 1 0 R >>\n".encode()
            + b"startxref\n" + str(xref).encode() + b"\n%%EOF")
    with open(path, "wb") as f:
        f.write(out)


_make_pdf(
    [
        [
            "Research Notes",
            "",
            "Chapter 1: Overview",
            "",
            "This is the first paragraph of the overview with enough text to act as a real",
            "block of body content used for testing the chunker and citations.",
            "",
            "Part of the methodology relied on careful sampling, and this sentence is a normal",
            "paragraph that must not become a heading even though it starts with the word Part.",
        ],
        [
            "Chapter 3 covered how the earlier sections connect, but this is body text and must",
            "remain a paragraph rather than a heading.",
            "",
            "A short closing line.",
        ],
    ],
    os.path.join(FIXDIR, "sample.pdf"),
)

print("fixtures written to", FIXDIR)
for fn in sorted(os.listdir(FIXDIR)):
    print("  ", fn, os.path.getsize(os.path.join(FIXDIR, fn)), "bytes")
