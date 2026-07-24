import unittest
from io import BytesIO

from docx import Document

from document_pipeline import chunker, document_map, ingest, pdf_layout, profiling


class DocumentIntelligenceTests(unittest.TestCase):
    def test_chaptered_document_builds_stable_parent_units(self):
        markdown = (
            "# کتاب نمونه\n\n"
            "## فصل اول: شروع\n\n" + ("متن فصل اول. " * 80) + "\n\n"
            "## فصل دوم: ادامه\n\n" + ("متن فصل دوم. " * 80)
        )
        meta = {"page_count": None, "structure_confidence": "high"}
        profile = profiling.profile_document(markdown, meta, filename="کتاب نمونه.txt")
        doc_map = document_map.build_document_map(markdown, profile)
        chunks = chunker.parse_markdown_to_chunks(markdown)
        document_map.assign_chunks_to_units(chunks, doc_map)

        self.assertEqual(profile.document_type, "chaptered_book")
        self.assertEqual(profile.unit_strategy, "chapter")
        self.assertGreaterEqual(len(doc_map["units"]), 2)
        self.assertTrue(all(chunk.get("parent_id") for chunk in chunks))

    def test_repeated_pdf_chapter_headers_do_not_create_duplicate_units(self):
        markdown = (
            "<!-- page:1 -->\n\n## فصل اول: شروع\n\nمتن صفحه اول.\n\n"
            "<!-- page:2 -->\n\n## فصل اول: شروع\n\nمتن صفحه دوم.\n\n"
            "<!-- page:3 -->\n\n## فصل دوم: ادامه\n\nمتن صفحه سوم."
        )
        meta = {"page_count": 3, "structure_confidence": "high"}
        profile = profiling.profile_document(markdown, meta, filename="book.pdf")
        doc_map = document_map.build_document_map(markdown, profile)
        titles = [unit["title"] for unit in doc_map["units"]]
        self.assertEqual(titles, ["فصل اول: شروع", "فصل دوم: ادامه"])

    def test_flat_document_uses_semantic_windows_without_llm(self):
        paragraphs = [f"این پاراگراف شماره {index} است. " * 40 for index in range(12)]
        markdown = "\n\n".join(paragraphs)
        meta = {"page_count": None, "structure_confidence": "low"}
        profile = profiling.profile_document(markdown, meta, filename="notes.txt")
        doc_map = document_map.build_document_map(markdown, profile)

        self.assertEqual(profile.document_type, "flat_document")
        self.assertEqual(profile.unit_strategy, "semantic_window")
        self.assertGreater(len(doc_map["units"]), 1)
        self.assertTrue(all(unit["title"].startswith("بخش ") for unit in doc_map["units"]))

    def test_unusable_ocr_output_is_rejected(self):
        markdown = "<!-- page:1 -->\n\n???"
        meta = {
            "page_count": 1,
            "structure_confidence": "low",
            "ocr_required": True,
            "ocr_status": "unavailable",
        }
        profile = profiling.profile_document(markdown, meta, filename="scan.pdf")

        self.assertEqual(profile.document_type, "noisy_scan")
        self.assertEqual(profile.quality.status, "rejected")
        self.assertFalse(profile.quality.indexable)

    def test_english_prose_is_not_misclassified_as_persian_mojibake(self):
        text = (
            "Artificial intelligence systems should be governed and evaluated "
            "throughout their lifecycle. " * 20
        )

        self.assertFalse(ingest._looks_garbled(text))

    def test_sparse_unavailable_ocr_pages_only_warn(self):
        markdown = "\n\n".join(
            [f"<!-- page:{page} -->\n\nمتن سالم صفحه {page}. " * 20 for page in range(1, 101)]
        )
        meta = {
            "page_count": 100,
            "structure_confidence": "high",
            "ocr_required": True,
            "ocr_status": "unavailable",
            "ocr_pages_skipped": 2,
            "ocr_blocking": False,
        }

        profile = profiling.profile_document(markdown, meta, filename="book.pdf")

        self.assertTrue(profile.quality.indexable)
        self.assertEqual(profile.quality.status, "warning")
        self.assertIn("partial_ocr_pages_skipped", profile.quality.warnings)

    def test_formal_meeting_speakers_are_detected_as_transcript(self):
        names = ["ADAMS", "BAKER", "CLARK", "DAVIS", "EVANS", "FRANK", "GREEN", "HALL", "IRWIN"]
        turns = [
            f"MR. {name}. This is a substantive meeting statement number {index}."
            for index, name in enumerate(names, start=1)
        ]
        markdown = "\n\n".join(turns)
        profile = profiling.profile_document(
            markdown,
            {"structure_confidence": "low"},
            filename="meeting.pdf",
        )

        self.assertEqual(profile.document_type, "transcript")
        self.assertEqual(profile.unit_strategy, "semantic_window")

    def test_numeric_table_row_is_not_promoted_to_heading(self):
        row = "62 Serbia 0.833 76.8 15.0 11.6 23,115 7 61"

        self.assertFalse(ingest._generic_heading_candidate(row))

    def test_docx_tables_are_preserved_in_reading_order(self):
        document = Document()
        document.add_heading("Report title", level=1)
        document.add_paragraph("Text before the table.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Year"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "2026"
        table.cell(1, 1).text = "42"
        document.add_paragraph("Text after the table.")
        stream = BytesIO()
        document.save(stream)
        stream.seek(0)

        result = ingest.normalize_document("report.docx", stream, ".docx")

        self.assertTrue(result["meta"]["has_tables"])
        self.assertEqual(result["meta"]["table_count"], 1)
        self.assertIn("| Year | Value |", result["markdown_text"])
        self.assertLess(result["markdown_text"].index("Text before"), result["markdown_text"].index("| Year"))
        self.assertLess(result["markdown_text"].index("| Year"), result["markdown_text"].index("Text after"))

    def test_broken_persian_pdf_spacing_is_repaired_without_rewriting(self):
        repaired = pdf_layout.normalize_visible("بررس ی رو یکرد د یوی د بوهم به عل یت")
        self.assertEqual(repaired, "بررسی رویکرد دیوید بوهم به علیت")

    def test_article_administration_is_classified_but_not_substantive(self):
        markdown = (
            "<!-- page:1 -->\n\n# عنوان مقاله\n\n## چکیده\n\nهدف پژوهش.\n\n"
            "<!-- page:2 -->\n\n## روش کار\n\nروش تاپسیس.\n\n"
            "## یافته ها\n\nنتایج اصلی.\n\n## نتیجه‌گیری\n\nپیام اصلی.\n\n"
            "## قدردانی ها\n\nاز حامیان سپاسگزاری شد.\n\n"
            "## مشارکت پدیدآوران\n\nنویسنده همه مراحل را انجام داد.\n\n"
            "## منابع مالی\n\nحمایت مالی اعلام شد.\n\n"
            "## ملاحظات اخلاقی\n\nکد اخلاق ثبت شد.\n\n"
            "## تعارض منافع\n\nتعارضی گزارش نشد.\n\n"
            "## References\n\nReference one."
        )
        profile = profiling.profile_document(
            markdown,
            {"page_count": 2, "document_title": "عنوان مقاله"},
            filename="paper.pdf",
        )
        doc_map = document_map.build_document_map(markdown, profile)
        roles = {unit["title"]: unit["role"] for unit in doc_map["units"]}

        self.assertEqual(profile.title, "عنوان مقاله")
        self.assertEqual(profile.document_type, "research_article")
        for title in ("قدردانی ها", "مشارکت پدیدآوران", "منابع مالی", "ملاحظات اخلاقی", "تعارض منافع"):
            self.assertEqual(roles[title], "administrative")
            self.assertFalse(document_map.is_substantive_section(title, roles[title]))
        self.assertEqual(roles["References"], "references")
        self.assertTrue(document_map.is_substantive_section("روش کار", roles["روش کار"]))
        self.assertTrue(document_map.is_substantive_section("یافته ها", roles["یافته ها"]))


if __name__ == "__main__":
    unittest.main()
